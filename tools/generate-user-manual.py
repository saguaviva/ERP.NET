from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "src" / "Erp.SaaS" / "docs"
DOCX_PATH = DOCS / "manual-usuario-coreflow.docx"
MD_PATH = DOCS / "manual-usuario-coreflow.md"
DOC_DATE = "24/04/2026"
VERSION = "v1.0"

POLISH_REPLACEMENTS = [
    ("Version", "Versión"),
    ("Indice", "Índice"),
    ("Como", "Cómo"),
    ("Donde", "Dónde"),
    ("Que", "Qué"),
    ("Guia", "Guía"),
    ("guia", "guía"),
    ("practica", "práctica"),
    ("practicas", "prácticas"),
    ("practico", "práctico"),
    ("rapida", "rápida"),
    ("rapidas", "rápidas"),
    ("rapido", "rápido"),
    ("rapidos", "rápidos"),
    ("aplicacion", "aplicación"),
    ("aplicaciones", "aplicaciones"),
    ("navegacion", "navegación"),
    ("catalogo", "catálogo"),
    ("parametros", "parámetros"),
    ("areas", "áreas"),
    ("area", "área"),
    ("Articulos", "Artículos"),
    ("articulos", "artículos"),
    ("articulo", "artículo"),
    ("Produccion", "Producción"),
    ("produccion", "producción"),
    ("Almacen", "Almacén"),
    ("almacen", "almacén"),
    ("Estadisticas", "Estadísticas"),
    ("estadisticas", "estadísticas"),
    ("Administracion", "Administración"),
    ("administracion", "administración"),
    ("funcion", "función"),
    ("accion", "acción"),
    ("Accion", "Acción"),
    ("funcional", "funcional"),
    ("vision", "visión"),
    ("historico", "histórico"),
    ("historica", "histórica"),
    ("logica", "lógica"),
    ("Codigo", "Código"),
    ("codigo", "código"),
    ("tecnica", "técnica"),
    ("tecnico", "técnico"),
    ("tecnicos", "técnicos"),
    ("composicion", "composición"),
    ("reposicion", "reposición"),
    ("valoracion", "valoración"),
    ("edicion", "edición"),
    ("revision", "revisión"),
    ("sesion", "sesión"),
    ("opcion", "opción"),
    ("lineas", "líneas"),
    ("pais", "país"),
    ("numero", "número"),
    ("Sintomas", "Síntomas"),
    ("criticas", "críticas"),
    ("telefono", "teléfono"),
    ("minimos", "mínimos"),
    ("minimas", "mínimas"),
    ("informacion", "información"),
    ("evolucion", "evolución"),
    ("concentracion", "concentración"),
    ("titulo", "título"),
    ("descripcion", "descripción"),
    ("paises", "países"),
    ("Paginacion", "Paginación"),
    ("pagina", "página"),
    ("aun", "aún"),
    ("dia", "día"),
    ("deberia", "debería"),
    ("gestion", "gestión"),
    ("autorizacion", "autorización"),
    ("correccion", "corrección"),
    ("desviacion", "desviación"),
    ("desviaciones", "desviaciones"),
    ("generico", "genérico"),
    ("busqueda", "búsqueda"),
    ("busquedas", "búsquedas"),
    ("ordenacion", "ordenación"),
    ("paginacion", "paginación"),
    ("despues", "después"),
    ("criticos", "críticos"),
    ("analisis", "análisis"),
    ("sincronizacion", "sincronización"),
    ("validacion", "validación"),
    ("previsualizacion", "previsualización"),
    ("exportacion", "exportación"),
    ("impresion", "impresión"),
    ("maquinas", "máquinas"),
    ("maquina", "máquina"),
    ("mercancia", "mercancía"),
    ("recepcion", "recepción"),
    ("albaran", "albarán"),
    ("Albaran", "Albarán"),
    ("albaranes", "albaranes"),
    ("Albaranes", "Albaranes"),
    ("facturacion", "facturación"),
    ("fabricacion", "fabricación"),
    ("formacion", "formación"),
    ("Formacion", "Formación"),
    ("creacion", "creación"),
    ("disposicion", "disposición"),
    ("especifico", "específico"),
    ("especificos", "específicos"),
    ("especifica", "específica"),
    ("especificas", "específicas"),
    ("ingles", "inglés"),
    ("menus", "menús"),
    ("menu", "menú"),
    ("segun", "según"),
    ("unicamente", "únicamente"),
    ("mas", "más"),
    ("ordenes", "órdenes"),
    ("Terminos", "Términos"),
    ("terminos", "términos"),
    ("categoria", "categoría"),
    ("categorias", "categorías"),
    ("configuracion", "configuración"),
    ("rapidas", "rápidas"),
    ("Direccion", "Dirección"),
    ("direccion", "dirección"),
    ("Resolucion", "Resolución"),
    ("resolucion", "resolución"),
    ("Sintoma", "Síntoma"),
    ("sintoma", "síntoma"),
    ("Validacion", "Validación"),
    ("Interpretacion", "Interpretación"),
    ("accion", "acción"),
    ("utiles", "útiles"),
    ("fisicas", "físicas"),
    ("codificacion", "codificación"),
    ("rotacion", "rotación"),
    ("antiguedad", "antigüedad"),
    ("modulo", "módulo"),
    ("modulos", "módulos"),
    ("regularizacion", "regularización"),
    ("Regularizaciones", "Regularizaciones"),
    ("regularizaciones", "regularizaciones"),
    ("vacio", "vacío"),
    ("automatica", "automática"),
    ("numeros", "números"),
    ("que esta pendiente", "qué está pendiente"),
    ("esta disponible", "está disponible"),
    ("esta en fase", "está en fase"),
    ("esta vivo", "está vivo"),
    ("este disponible", "esté disponible"),
    ("donde resolverlas", "dónde resolverlas"),
    ("por donde", "por dónde"),
    ("Para que sirve", "Para qué sirve"),
]


def polish_text(text: str) -> str:
    for source, target in POLISH_REPLACEMENTS:
        text = re.sub(rf"(?<![A-Za-zÀ-ÿ]){re.escape(source)}(?![A-Za-zÀ-ÿ])", target, text)
    return text


MANUAL_SECTIONS = [
    ("Primeros pasos", "Preparar el entorno de trabajo antes de entrar en procesos operativos."),
    ("Guia rapida para usuarios", "Accesos directos a las tareas mas habituales y donde resolverlas."),
    ("Rutinas por rol", "Que deberia revisar cada perfil al empezar, durante el dia y antes de cerrar."),
    ("Como leer una pantalla", "Partes comunes de una pantalla: cabecera, filtros, KPIs, tabla, acciones y documentos."),
    ("Capturas recomendadas", "Pantallas que conviene capturar para convertir el manual en una guia visual de formacion."),
    ("Estados, señales y datos criticos", "Interpretar estados, avisos y campos que no conviene dejar sin revisar."),
    ("Conceptos comunes", "Empresa activa, idioma, idioma documental, tema visual, navegacion, filtros y documentos."),
    ("Mapa funcional", "Resumen de las areas disponibles y pantallas principales."),
    ("Workflows transversales", "Operaciones repetidas en toda la aplicacion: buscar, crear, editar, imprimir y exportar."),
    ("CRM", "Clientes, proveedores, duplicados, representantes, transportistas y talleres."),
    ("Articulos", "Fornituras, hilos, tejidos, models, muestras, complementos y disposiciones."),
    ("Base", "Maestros y parametros que alimentan compras, ventas, produccion y almacen."),
    ("Produccion", "Partes de acabado y ordenes de fabricacion asociadas a muestras, models y disposiciones."),
    ("Ventas", "Pedidos, albaranes, pre-facturacion, facturas, remesas e Intrastat."),
    ("Compras", "Pedidos, recepciones, facturas de proveedor y circuitos especificos de hilos y fabricacion."),
    ("Almacen", "Stock, traspasos, movimientos, inventarios, conteo ciego y ajustes."),
    ("Workflows operativos principales", "Procedimientos concretos para resolver tareas habituales de trabajo."),
    ("Procesos punta a punta", "Secuencias completas desde el dato maestro hasta documento, produccion, stock o reporting."),
    ("Criterios de calidad", "Que comprobar antes de guardar, imprimir, exportar o cerrar un proceso."),
    ("Resolucion de incidencias", "Sintomas habituales, causa probable y primera accion recomendada."),
    ("Checklists de cierre", "Rutinas de cierre diario, semanal y mensual para evitar trabajo pendiente oculto."),
    ("Listados y estadisticas", "Consultas ejecutivas, reporting por areas, exportaciones y lectura de KPIs."),
    ("Alertas operativas", "Avisos de trabajo pendiente y señales de datos que conviene revisar."),
    ("Administracion", "Empresas, usuarios, tenant, plataforma y configuracion de acceso."),
    ("Buenas practicas", "Recomendaciones para trabajar con seguridad, consistencia y trazabilidad."),
    ("Preguntas frecuentes", "Respuestas rapidas a incidencias habituales."),
    ("Glosario", "Terminos frecuentes usados en el ERP."),
]


FUNCTIONAL_MAP = [
    ("Dashboard", "Vision inicial por areas y accesos rapidos.", "Inicio"),
    ("CRM", "Clientes, duplicados, proveedores, representantes, transportistas y talleres.", "/crm/clientes, /crm/proveedores, /crm/talleres"),
    ("Articulos", "Catalogo textil y de producto: fornituras, hilos, tejidos, models, muestras, complementos y disposiciones.", "/articulos/tejidos, /articulos/muestras, /articulos/models"),
    ("Base", "Maestros comunes: numeraciones, bancos, formas de pago, operaciones, maquinas, temporadas, IVA e incoterms.", "/base-datos/..."),
    ("Produccion", "Partes de acabado, ordenes de fabricacion y seguimiento de estado.", "/produccion/acabados"),
    ("Ventas", "Pedidos, albaranes, pre-facturacion, facturas, remesas e Intrastat.", "/ventas/..."),
    ("Compras", "Proveedores, pedidos, recepciones, facturas y órdenes específicas.", "/compras/..."),
    ("Almacen", "Stock actual, hilos, tejidos, models, movimientos, traspasos, inventarios y ajustes.", "/almacen/..."),
    ("Listados", "Cuadros y consultas operativas por ventas, compras, produccion y almacen.", "/listados"),
    ("Estadisticas", "KPIs, comparativas, evolucion y analitica por areas.", "/estadisticas"),
    ("Alertas", "Panel de avisos operativos y revision de incidencias.", "/alertas"),
    ("Administracion", "Gestion de tenant, plataforma, empresas y usuarios con permiso.", "/administracion, /plataforma"),
]


CONCEPTS = [
    ("Empresa activa", "Todos los listados y documentos se consultan dentro de la empresa seleccionada. Cambiar empresa debe mantener la pantalla y recargar datos."),
    ("Idioma de la web", "Afecta a menus, botones, etiquetas y textos de pantalla."),
    ("Idioma de documentos", "Afecta a PDFs, impresiones y exportaciones. Puede ser distinto del idioma de la web."),
    ("Tema visual", "Cambia densidad y aspecto. Ultracompacto aprovecha casi toda la pantalla dejando margen de lectura."),
    ("Menu lateral", "Organizado por categorias plegables: CRM, Articulos, Base, Produccion, Ventas, Compras, Almacen y Administracion."),
    ("Origen legacy", "Indica que el registro procede del sistema historico. La web guarda cambios locales salvo procesos de sincronizacion definidos."),
    ("Listados", "Normalmente permiten busqueda, filtros, ordenacion, paginacion y acceso a ficha o impresion."),
    ("Documentos", "Pedidos, albaranes, facturas, partes, inventarios y remesas conservan trazabilidad entre pasos."),
]


QUICK_RECIPES = [
    ("Quiero encontrar un cliente", "CRM > Clientes", "Buscar por nombre, NIF, ciudad, email o telefono. Revisar duplicados si hay varias coincidencias."),
    ("Quiero crear un cliente o proveedor", "CRM > Clientes / Proveedores", "Buscar antes de crear. Completar datos fiscales y de contacto minimos antes de usarlo en documentos."),
    ("Quiero consultar una ficha de tejido", "Articulos > Tejidos", "Buscar por codigo o descripcion. Abrir ficha, revisar datos tecnicos, colores, escandallo, stock e impresiones."),
    ("Quiero lanzar un trabajo a planta", "Produccion > Acabados", "Crear parte manual o desde muestra/disposicion. Revisar acabador, maquina, cantidades, estado e imprimir parte."),
    ("Quiero preparar un inventario", "Almacen > Inventarios", "Crear inventario, filtrar alcance, imprimir hoja de conteo ciego, introducir recuentos y validar diferencias."),
    ("Quiero revisar una factura de cliente", "Ventas > Facturas", "Buscar factura, abrir detalle, revisar cliente, albaranes, cobro, pendiente, contabilidad, total y PDF."),
    ("Quiero exportar informacion", "Listados / Estadisticas / Intrastat", "Aplicar filtros, comprobar idioma documental, exportar y validar totales antes de compartir."),
    ("Quiero saber que esta pendiente", "Alertas / Dashboard / Estadisticas", "Empezar por alertas, luego KPIs y finalmente listados concretos del area afectada."),
]


ROLE_ROUTINES = [
    ("Administracion", "Revisar alertas, facturas pendientes, remesas, formas de pago, cobros/pagos y documentos con datos fiscales incompletos."),
    ("Comercial", "Revisar clientes, pedidos, albaranes, facturas, duplicados y trazabilidad de documentos por cliente."),
    ("Compras", "Revisar proveedores, pedidos, recepciones, facturas proveedor, diferencias recepcion-factura y stock de hilos."),
    ("Produccion / Planta", "Revisar partes vivos, disposiciones, acabadores, maquinas, impresiones de parte y carga por estado/semana."),
    ("Almacen", "Revisar stock actual, movimientos, traspasos, inventarios abiertos, conteos ciegos y diferencias pendientes de validar."),
    ("Direccion", "Revisar estadisticas por area: ventas, compras, produccion, almacen, evolucion, concentracion y alertas principales."),
]


SCREEN_ANATOMY = [
    ("Cabecera", "Indica modulo, titulo, descripcion y acciones principales como Nuevo, Volver, PDF o Imprimir."),
    ("Empresa activa", "Determina los datos visibles. Si no ves registros esperados, comprueba primero la empresa."),
    ("Filtros", "Permiten acotar por texto, estado, fechas, almacen, proveedor, cliente, articulo o color segun pantalla."),
    ("KPIs / tarjetas", "Resumen rapido de volumen, importes, diferencias, lineas, paises, stock o actividad."),
    ("Tabla", "Lista los registros. Normalmente permite ordenar columnas y abrir detalle."),
    ("Acciones de fila", "Abrir, editar, imprimir, crear parte, PDF o acciones especificas del documento."),
    ("Paginacion", "Controla pagina y tamaño. Si faltan resultados, revisa filtros antes de aumentar pagina."),
    ("Mensajes", "Avisos de validacion, errores o notas de sincronizacion. No conviene ignorarlos."),
]


CAPTURE_PLAN = [
    ("Inicio / Dashboard", "Mostrar acceso inicial, tarjetas por area y navegacion principal.", "Usuario nuevo entiende por donde empezar."),
    ("Menu lateral plegado/desplegado", "Mostrar categorias CRM, Articulos, Base, Produccion, Ventas, Compras, Almacen y Administracion.", "Explicar que no faltan opciones: pueden estar plegadas."),
    ("Selector de empresa e idioma documental", "Mostrar empresa activa, idioma web e idioma de impresiones/exportaciones.", "Evitar errores por empresa o idioma antes de crear documentos."),
    ("CRM / Clientes", "Mostrar busqueda, filtros, tabla, estado y acceso a duplicados.", "Formacion de usuarios comerciales y administracion."),
    ("Articulos / Tejidos ficha", "Mostrar ficha tecnica, colores, costes, stock e impresiones.", "Explicar la ficha mas rica del catalogo textil."),
    ("Articulos / Muestras o Complementos", "Mostrar desglose tecnico y acciones de impresion/parte.", "Explicar el puente hacia produccion."),
    ("Produccion / Acabados", "Mostrar listado, estado, filtros e impresion de parte.", "Formacion de planta y responsable de produccion."),
    ("Ventas / Facturas", "Mostrar cliente, estado, pendiente, contabilidad, total y PDF.", "Explicar revision antes de enviar o contabilizar."),
    ("Almacen / Inventarios", "Mostrar creacion, conteo ciego, diferencias y cierre.", "Formacion de almacen y recuentos de planta."),
    ("Estadisticas", "Mostrar rango de fechas, KPIs, comparativa y exportacion.", "Formacion para direccion y seguimiento semanal."),
    ("Alertas", "Mostrar lista de avisos, prioridad y enlace a origen.", "Explicar rutina diaria de revision."),
]


STATUS_GUIDE = [
    ("Borrador", "Documento editable que aun no deberia considerarse definitivo."),
    ("Emitida / Preparada", "Documento ya generado o listo para contabilidad/gestion segun circuito."),
    ("Pendiente", "Queda accion por realizar: cobrar, pagar, recibir, validar o cerrar."),
    ("Cerrado", "Proceso finalizado. Solo modificar si hay autorizacion o correccion justificada."),
    ("Anulado", "Documento sin efecto operativo, pero mantenido por trazabilidad."),
    ("Con diferencia", "Hay desviacion entre esperado y real. Revisar antes de cerrar inventario, recepcion o ajuste."),
    ("Origen legacy", "Registro importado del sistema historico. Modificar con especial cuidado."),
    ("Origen local", "Registro creado o modificado en la SaaS."),
]


CRITICAL_DATA = [
    ("Cliente / proveedor", "Debe mostrar nombre reconocible, no solo codigo generico, antes de enviar documentos."),
    ("Empresa y centro", "Validar siempre antes de crear documentos o sincronizar datos."),
    ("Fechas", "Afectan a facturacion, Intrastat, estadisticas, recepciones e inventarios."),
    ("Codigo articulo / color", "Clave para stock, producción, escandallos y trazabilidad."),
    ("Cantidades e importes", "Revisar separadores, decimales, signos negativos y totales."),
    ("Estado", "Define si algo esta vivo, pendiente, cerrado, anulado o requiere accion."),
    ("Idioma documental", "Validar antes de generar PDF, impresion o exportacion para terceros."),
]


WORKFLOWS = [
    ("Cambiar empresa activa", [
        "Abrir el selector Empresa activa del menu lateral.",
        "Seleccionar la empresa de trabajo.",
        "La pantalla actual se mantiene y se recarga con datos de esa empresa.",
        "Verificar el nombre de empresa antes de crear documentos o movimientos.",
    ]),
    ("Cambiar idioma de la web", [
        "Abrir el selector Idioma.",
        "Elegir Castellano, Català o English.",
        "La navegacion se mantiene en la pantalla actual.",
        "Recordar que el idioma de pantalla no cambia necesariamente el idioma de PDFs o exportaciones.",
    ]),
    ("Cambiar idioma de impresiones y exportaciones", [
        "Abrir Idioma impresiones / exportaciones.",
        "Elegir el idioma documental deseado.",
        "Generar de nuevo el PDF, impresion o exportacion.",
        "Usar este selector cuando el usuario trabaja en un idioma pero el cliente necesita documentos en otro.",
    ]),
    ("Buscar y filtrar un listado", [
        "Entrar en el area correspondiente desde el menu lateral.",
        "Escribir el criterio de busqueda: codigo, cliente, articulo, documento, color o notas segun pantalla.",
        "Ajustar filtros adicionales como estado, fechas, almacen o solo clasificados.",
        "Pulsar Buscar y revisar el contador de resultados.",
        "Ordenar columnas pulsando sobre el encabezado cuando este disponible.",
    ]),
    ("Crear un registro maestro", [
        "Entrar en el listado del maestro.",
        "Pulsar Nuevo, Nueva o el boton equivalente.",
        "Rellenar los campos obligatorios y los datos operativos.",
        "Guardar y volver al listado para comprobar que aparece.",
        "Si el registro viene de legacy, revisar que el origen y centro sean correctos antes de modificarlo.",
    ]),
    ("Editar un registro existente", [
        "Buscar el registro en el listado.",
        "Pulsar Abrir o Editar.",
        "Modificar unicamente los campos necesarios.",
        "Guardar y revisar mensajes de validacion.",
        "En registros legacy, recordar que los cambios son locales en la SaaS salvo que se indique lo contrario.",
    ]),
    ("Imprimir o generar PDF", [
        "Abrir el documento o ficha.",
        "Comprobar primero el idioma documental.",
        "Pulsar Imprimir, PDF o Abrir ficha imprimible.",
        "Revisar la previsualizacion antes de enviarlo al cliente o a planta.",
        "Guardar el PDF con un nombre que incluya documento, cliente y fecha.",
    ]),
    ("Exportar datos", [
        "Aplicar filtros antes de exportar.",
        "Comprobar el idioma de impresiones / exportaciones.",
        "Usar Exportar CSV, Excel o el boton especifico de la pantalla.",
        "Abrir el archivo y revisar columnas, separadores de miles y decimales antes de compartir.",
    ]),
]


AREA_DETAILS = {
    "CRM": [
        ("Clientes", "Buscar, crear, editar y revisar clientes. La bandeja de duplicados ayuda a detectar registros repetidos o dudosos."),
        ("Duplicados", "Muestra candidatos a duplicado para limpiar el maestro sin perder trazabilidad."),
        ("Proveedores", "Aunque aparece dentro del circuito de compras, forma parte del dato maestro comercial."),
        ("Representantes", "Mantiene la red comercial asociable a clientes y documentos."),
        ("Transportistas", "Mantiene agencias y transportistas usados en albaranes y expediciones."),
        ("Talleres", "Maestro operativo que puede alimentar produccion, articulos y compras."),
    ],
    "Articulos": [
        ("Fornituras", "Consulta y mantenimiento de complementos de producto usados en fichas y costes."),
        ("Hilos", "Maestro de hilos con proveedor, coste, precio, IVA, notas y carta de colores."),
        ("Tejidos", "Ficha tecnica con composicion, gramaje, ancho, costes, stock, colores, escandallo, etiquetas, reposicion y valoracion."),
        ("Models", "Catalogo especifico de models del legacy, con edicion cuidadosa porque el origen historico puede estar incompleto."),
        ("Muestras", "Ficha de muestra con desglose tecnico, colores/tallas, impresion y conexion con produccion."),
        ("Complementos", "Circuito paralelo a muestras para articulos complementarios y desglose tecnico."),
        ("Disposiciones", "Documentos de preparacion/trabajo con cliente, acabador, color, piezas, kilos, impresion y creacion de parte."),
    ],
    "Base": [
        ("Empresas", "Gestion de empresas disponibles para el tenant cuando el usuario tiene permisos."),
        ("Numeraciones", "Control de series y contadores de documentos."),
        ("Mailing", "Campañas y comunicaciones imprimibles."),
        ("Bancos / cajas", "Maestro financiero usado en cobros, pagos y remesas."),
        ("Formas de pago", "Condiciones reutilizables en clientes, proveedores y documentos."),
        ("Operaciones", "Catalogo operativo para procesos productivos."),
        ("Maquinas", "Maestro de maquinas utilizado en produccion, muestras y acabados."),
        ("Temporadas", "Clasificacion comercial o de coleccion."),
        ("Tipos de IVA", "Fiscalidad aplicable a documentos y maestros."),
        ("Incoterms", "Condiciones internacionales para operaciones comerciales."),
    ],
    "Produccion": [
        ("Partes de acabado", "Registrar trabajo productivo, acabador, maquina, estado, origen, cantidades y notas."),
        ("Alta guiada", "Puede iniciarse desde una muestra, complemento o disposicion cuando el flujo lo permite."),
        ("Impresion", "Genera un documento de planta con la informacion necesaria para ejecutar el trabajo."),
        ("Seguimiento", "El listado permite filtrar trabajo vivo y revisar estados para priorizar planta."),
    ],
    "Ventas": [
        ("Pedidos Teixits-Mostres", "Circuito de pedidos especifico para tejidos y muestras."),
        ("Albaranes", "Documentos de entrega y expedicion, con variantes para tejidos/muestras y models."),
        ("Pre-facturacion", "Borradores previos a factura para revisar importes y trazabilidad."),
        ("Facturas", "Facturas emitidas con PDF, estado de cobro, contabilidad, albaranes asociados e importe pendiente."),
        ("Remesas", "Agrupacion de cobros para gestion bancaria."),
        ("Intrastat", "Generacion y revision de datos comunitarios por mes, año, pais y codigo NC."),
    ],
    "Compras": [
        ("Pedidos", "Pedidos generales de compra."),
        ("Comandes fils", "Pedidos especificos de hilos."),
        ("Ordenes de fabricacion muestras/models", "Ordenes productivas ubicadas historicamente en compras dentro del legacy."),
        ("Recepciones", "Entrada de mercancia y control frente a pedido."),
        ("Albaranes entrega fils", "Recepciones especificas del circuito de hilos."),
        ("Facturas proveedor", "Registro y revision de facturas recibidas."),
    ],
    "Almacen": [
        ("Stock actual", "Vista general de existencias por articulo, color, almacen y posicion."),
        ("Stock / Fils", "Consulta operativa de hilos con filtros por proveedor, color y tipo de movimiento."),
        ("Stock / Tejidos", "Vista especifica de tejidos para cerrar simetria con el legacy."),
        ("Stock / Models", "Vista especifica de stock asociado a models."),
        ("Inventarios", "Conteos por almacen, familia, articulo o color; incluye conteo ciego e impresion para planta."),
        ("Traspasos", "Movimientos entre almacenes o ubicaciones."),
        ("Movimientos", "Historico de entradas, salidas, ajustes y actividad por posicion."),
        ("Ajustes", "Regularizaciones controladas cuando hay diferencias justificadas."),
    ],
}


AREA_PRACTICES = {
    "CRM": [
        "Buscar por NIF, email o telefono antes de crear un nuevo cliente o proveedor.",
        "Revisar la bandeja de duplicados cuando aparezcan nombres parecidos.",
        "Mantener datos fiscales y de contacto claros para evitar PDFs incompletos.",
        "No bloquear ni reactivar clientes sin confirmar el motivo con administracion.",
    ],
    "Articulos": [
        "Buscar por codigo legacy y descripcion antes de crear una ficha nueva.",
        "Completar datos tecnicos minimos antes de usar el articulo en ventas, compras o produccion.",
        "Revisar colores, composicion, costes y stock antes de imprimir fichas o etiquetas.",
        "En models, ser especialmente prudente: el origen historico puede estar incompleto.",
    ],
    "Base": [
        "Cambiar maestros solo cuando se entienda donde se usan.",
        "Evitar duplicar formas de pago, maquinas u operaciones con nombres ligeramente distintos.",
        "Revisar numeraciones antes de empezar una serie documental nueva.",
        "Mantener incoterms, IVA y bancos/cajas con nombres comprensibles para usuario final.",
    ],
    "Produccion": [
        "No lanzar partes sin acabador, maquina, cantidades y notas suficientes para planta.",
        "Usar alta guiada desde muestra o disposicion cuando exista, porque conserva trazabilidad.",
        "Imprimir el parte antes de pasar trabajo a planta.",
        "Revisar partes antiguos o bloqueados antes de crear nuevos.",
    ],
    "Ventas": [
        "Antes de facturar, revisar cliente, albaranes asociados, importes y estado de cobro.",
        "Usar el idioma documental correcto antes de generar PDF.",
        "En Teixits-Mostres, seguir el circuito pedido > albaran > pre-facturacion > factura.",
        "En Intrastat, comprobar empresa, mes, año, pais y codigo NC antes de exportar.",
    ],
    "Compras": [
        "Cruzar pedido, recepcion y factura antes de cerrar la revision del proveedor.",
        "En hilos, revisar proveedor, color y cantidad recibida contra stock.",
        "No registrar factura de proveedor sin validar diferencias relevantes de recepcion.",
        "Mantener proveedores completos para evitar incidencias fiscales o de trazabilidad.",
    ],
    "Almacen": [
        "Consultar movimientos antes de hacer ajustes manuales.",
        "Usar traspasos para mover stock entre almacenes o posiciones, no ajustes si hay traslado real.",
        "En inventarios, usar conteo ciego para planta y mostrar diferencias solo al validar.",
        "Cerrar inventarios solo cuando las diferencias esten revisadas por responsable.",
    ],
}


SPECIFIC_WORKFLOWS = [
    ("Resolver duplicados de cliente", [
        "Entrar en CRM > Clientes > Bandeja de duplicados.",
        "Filtrar por estado o buscar por nombre, NIF, ciudad, email o telefono.",
        "Comparar candidatos antes de actuar: nombre fiscal, NIF, email, telefono, ciudad y origen.",
        "Marcar para revisar cuando no haya seguridad suficiente.",
        "Evitar eliminar datos sin validacion comercial o administrativa.",
    ]),
    ("Crear o revisar un tejido", [
        "Entrar en Articulos > Tejidos.",
        "Buscar por codigo, descripcion, ancho, maquina, tejedor o acabador.",
        "Abrir la ficha o crear un nuevo tejido.",
        "Completar datos tecnicos: composicion, gramaje, ancho, tubular, costes y notas.",
        "Revisar carta de colores, escandallo, acabados y stock.",
        "Usar las impresiones de ficha, etiquetas, reposicion o valoracion cuando proceda.",
    ]),
    ("Pasar de muestra a produccion", [
        "Entrar en Articulos > Muestras o Complementos.",
        "Abrir la ficha y revisar el desglose tecnico concreto.",
        "Seleccionar la accion de crear parte u orden de fabricacion cuando este disponible.",
        "Comprobar color, maquina, cantidades, acabador y notas antes de guardar.",
        "Imprimir el parte para planta si el trabajo ya puede lanzarse.",
    ]),
    ("Trabajar con disposiciones", [
        "Entrar en Articulos > Disposiciones.",
        "Filtrar por numero, cliente, acabador, color, comanda o notas.",
        "Abrir la disposicion para revisar piezas, kilos, estado y origen.",
        "Imprimir el documento o crear parte de acabado cuando proceda.",
        "Evitar crear parte si faltan datos criticos de cliente, color o acabador.",
    ]),
    ("Circuito de ventas Teixits-Mostres", [
        "Entrar en Ventas > Pedidos Teixits-Mostres.",
        "Registrar o revisar pedido y lineas.",
        "Generar o consultar albaranes asociados cuando hay entrega.",
        "Revisar pre-facturacion antes de emitir factura.",
        "Abrir la factura, comprobar pendiente, cobro, contabilidad y total.",
        "Generar PDF en el idioma documental correcto.",
    ]),
    ("Circuito de compras de hilos", [
        "Entrar en Compras > Comandes fils.",
        "Registrar pedido al proveedor con articulo/color y cantidades.",
        "Al recibir mercancia, consultar o registrar albaran de entrega fils.",
        "Revisar recepcion frente a factura para detectar diferencias.",
        "Consultar Stock / Fils para validar impacto en almacen.",
    ]),
    ("Revisar factura antes de enviarla o contabilizarla", [
        "Abrir Ventas > Facturas o Compras > Facturas segun corresponda.",
        "Buscar por numero, cliente/proveedor, fecha o notas.",
        "Comprobar nombre fiscal, fecha, vencimiento, lineas, base, IVA, total y estado.",
        "Verificar si hay albaranes o recepciones asociados.",
        "Generar PDF solo despues de comprobar el idioma documental.",
    ]),
    ("Preparar Intrastat mensual", [
        "Entrar en Ventas > Intrastat.",
        "Seleccionar empresa, mes y año.",
        "Revisar paises UE, lineas clasificadas, lineas sin NC code y base Intrastat.",
        "Abrir detalle si hay importes inesperados, negativos o paises no previstos.",
        "Exportar Excel/CSV y validar totales antes de presentar o compartir.",
    ]),
    ("Hacer un traspaso de stock", [
        "Entrar en Almacen > Traspasos.",
        "Crear nuevo traspaso.",
        "Indicar origen, destino, articulo, color y cantidad.",
        "Comprobar que el movimiento representa un traslado real y no una regularizacion.",
        "Guardar y revisar Stock actual o Movimientos para confirmar el impacto.",
    ]),
    ("Inventario con conteo ciego", [
        "Entrar en Almacen > Inventarios.",
        "Crear inventario nuevo o nuevo desde stock.",
        "Filtrar por almacen, familia, articulo o color si el conteo es parcial.",
        "Imprimir hoja de recuento ciego para planta: articulo, color y casilla de conteo, sin stock esperado ni importes.",
        "Introducir conteos reales.",
        "Validar diferencias de forma masiva por familia o almacen.",
        "Cerrar el inventario solo cuando las diferencias esten revisadas.",
    ]),
    ("Consultar estadisticas por area", [
        "Entrar en Estadisticas.",
        "Seleccionar rango de fechas.",
        "Abrir Ventas, Compras, Produccion o Almacen.",
        "Leer primero KPIs principales, luego comparativa y evolucion.",
        "Exportar CSV si se necesita analisis externo.",
    ]),
    ("Revisar alertas operativas", [
        "Entrar en Alertas.",
        "Priorizar avisos criticos o con impacto en cliente, stock o produccion.",
        "Abrir el enlace asociado al documento o maestro.",
        "Corregir el dato origen si procede.",
        "Volver a alertas para comprobar si el aviso desaparece o queda pendiente.",
    ]),
    ("Cierre operativo semanal", [
        "Revisar Alertas y resolver incidencias criticas.",
        "Consultar Estadisticas por area para detectar desviaciones.",
        "Revisar facturas pendientes, recepciones sin factura y partes de produccion vivos.",
        "Comprobar inventarios abiertos y diferencias no validadas.",
        "Exportar listados necesarios y guardar evidencias antes de cerrar el periodo.",
    ]),
]


END_TO_END_FLOWS = [
    (
        "Alta de cliente hasta factura",
        "Cliente o pedido recibido",
        "CRM > Cliente; Ventas > Pedido; Albaran; Pre-facturacion; Factura",
        "Factura emitida con cliente, lineas, total y PDF revisado",
        "Nombre fiscal correcto, importes coherentes, idioma documental adecuado y estado de cobro claro.",
    ),
    (
        "Muestra hasta parte de acabado",
        "Muestra o complemento con desglose tecnico",
        "Abrir ficha; revisar desglose; crear parte; asignar acabador/maquina; imprimir",
        "Parte listo para planta con trazabilidad a origen",
        "Color, cantidad, maquina, acabador y notas de planta completos.",
    ),
    (
        "Compra de hilos hasta stock",
        "Necesidad de hilo o pedido a proveedor",
        "Compras > Comandes fils; recepcion; factura proveedor; Almacen > Stock / Fils",
        "Stock de hilo actualizado y recepcion trazable",
        "Proveedor, color, cantidad recibida y factura cruzados sin diferencias relevantes.",
    ),
    (
        "Disposicion hasta produccion",
        "Disposicion preparada para trabajar",
        "Articulos > Disposiciones; abrir; imprimir o crear parte de acabado",
        "Trabajo preparado para planta/acabador",
        "Cliente, color, piezas, kilos, estado y acabador revisados.",
    ),
    (
        "Inventario hasta ajuste",
        "Necesidad de validar stock real",
        "Almacen > Inventarios; conteo ciego; introducir recuento; validar diferencias; cerrar",
        "Stock regularizado con diferencias revisadas",
        "Diferencias justificadas por familia/almacen y cierre autorizado.",
    ),
    (
        "Intrastat mensual",
        "Periodo mensual cerrado o en revision",
        "Ventas > Intrastat; seleccionar empresa/mes/año; revisar detalle; exportar",
        "Archivo de Intrastat preparado para revision externa",
        "Paises UE, NC code, importes y negativos revisados antes de presentar.",
    ),
]


QUALITY_GATES = [
    ("Antes de crear", "Buscar duplicados y confirmar empresa activa.", "Evitar registros repetidos y datos en empresa equivocada."),
    ("Antes de guardar maestro", "Codigo, nombre, descripcion, origen, notas y campos obligatorios.", "Registro localizable y reutilizable en otros modulos."),
    ("Antes de guardar documento", "Cliente/proveedor, fecha, lineas, cantidades, precios, estado y observaciones.", "Documento coherente y trazable."),
    ("Antes de imprimir/PDF", "Idioma documental, destinatario, importes, datos fiscales y logo/encabezado.", "Documento listo para cliente, proveedor o planta."),
    ("Antes de exportar", "Filtros, rango de fechas, empresa, totales y formato numerico.", "Archivo consistente para Excel o reporting externo."),
    ("Antes de cerrar inventario", "Diferencias, responsable, familias afectadas y movimientos pendientes.", "Cierre sin regularizaciones ocultas."),
    ("Antes de cerrar semana", "Alertas, facturas pendientes, recepciones sin factura, partes vivos e inventarios abiertos.", "Periodo operativo controlado."),
]


INCIDENT_PLAYBOOK = [
    ("No aparecen datos", "Empresa, filtros, fechas o permisos no coinciden.", "Limpiar filtros, comprobar empresa activa, ampliar fechas y revisar rol."),
    ("Aparece Cliente 123 en vez del nombre", "La consulta no ha resuelto el maestro de cliente.", "Abrir cliente por codigo y revisar que exista nombre; si se repite, apuntar pantalla para correccion."),
    ("El PDF sale en otro idioma", "Se cambió idioma web pero no idioma documental.", "Cambiar Idioma impresiones / exportaciones y generar de nuevo."),
    ("Un total parece incorrecto", "Filtro parcial, linea negativa, abono o formato de decimales.", "Abrir detalle, revisar lineas, signos, base, IVA y total."),
    ("No veo una opcion del menu", "Categoria plegada o falta de permisos.", "Desplegar categoria; si no aparece, consultar rol con administracion."),
    ("No se puede cerrar inventario", "Diferencias sin validar o lineas incompletas.", "Filtrar por diferencias, revisar conteos y validar masivamente solo cuando proceda."),
    ("El listado va muy cargado", "Rango o filtro demasiado amplio.", "Usar busqueda, fechas y tamaño de pagina menor; en ultracompacto ocultar menu lateral."),
    ("Intrastat no cuadra", "Mes/empresa incorrectos, NC code faltante o pais no UE.", "Revisar detalle por factura y agrupacion por pais/NC code antes de exportar."),
]


CLOSING_CHECKLISTS = {
    "Cierre diario": [
        "Revisar Alertas y resolver las criticas.",
        "Comprobar documentos creados durante el dia: cliente/proveedor, fecha, importes y estado.",
        "Revisar partes de produccion vivos que deban imprimirse o asignarse.",
        "Comprobar movimientos de almacen relevantes y traspasos pendientes.",
        "Anotar incidencias que no puedan resolverse en el dia.",
    ],
    "Cierre semanal": [
        "Revisar Estadisticas por area y comparar con el periodo anterior.",
        "Cruzar recepciones de compras con facturas proveedor.",
        "Revisar facturas de venta pendientes de cobro o contabilizacion.",
        "Revisar inventarios abiertos y diferencias sin validar.",
        "Exportar listados necesarios para seguimiento interno.",
    ],
    "Cierre mensual": [
        "Revisar Intrastat del mes por empresa antes de exportar.",
        "Comprobar remesas, cobros, pagos y facturas pendientes.",
        "Revisar stock parado, rotacion y diferencias de inventario.",
        "Validar que no queden partes productivos antiguos sin estado claro.",
        "Guardar exportaciones y evidencias en la carpeta de cierre correspondiente.",
    ],
}


BEST_PRACTICES = [
    "Comprueba siempre la empresa activa antes de crear documentos, movimientos o fichas.",
    "Usa busquedas y filtros antes de crear un nuevo registro para evitar duplicados.",
    "Distingue entre idioma de pantalla e idioma de documentos: pueden ser distintos.",
    "No modifiques datos legacy sin revisar primero origen, centro y trazabilidad.",
    "En listados grandes, guarda o comparte la URL cuando quieras conservar filtros y paginacion.",
    "Antes de cerrar inventarios, revisa diferencias por familia o almacen y documenta incidencias.",
    "Antes de enviar PDFs a cliente o proveedor, abre la previsualizacion y valida importes, idioma y datos fiscales.",
    "Usa alertas y estadisticas como rutina diaria: primero avisos, despues trabajo operativo.",
    "En produccion, no lances partes sin acabador, maquina, cantidades y notas minimas de planta.",
    "En exportaciones, verifica separadores, totales y filtros aplicados antes de enviar a terceros.",
]


FAQ = [
    ("Cambio de idioma y vuelvo a otra pantalla", "El comportamiento esperado es permanecer en la pantalla actual. Si ocurre, revisar si la URL tiene parametros raros o si la sesion ha caducado."),
    ("Cambio de empresa y no veo datos", "Puede que la empresa activa no tenga registros para ese modulo o que el usuario no tenga permisos sobre esa empresa."),
    ("Veo 'Cliente 123' en vez del nombre", "Significa que el documento tiene codigo de cliente pero no se ha resuelto el nombre en esa consulta. Conviene revisar el maestro de clientes y la consulta de esa pantalla."),
    ("Un listado parece vacio", "Revisar filtros activos, empresa, fechas, estado y tamaño de pagina. En Intrastat, comprobar mes/año y empresa."),
    ("No aparece una opcion de menu", "Puede depender del rol del usuario o estar dentro de una categoria plegada del menu lateral."),
    ("El PDF sale en otro idioma", "Revisar el selector Idioma impresiones / exportaciones, no solo el idioma de la web."),
]


GLOSSARY = [
    ("Tenant", "Grupo o entorno SaaS que contiene empresas, usuarios y configuracion compartida."),
    ("Empresa activa", "Empresa sobre la que se filtran listados, documentos y maestros operativos."),
    ("Legacy", "Sistema historico original desde el que se importan o sincronizan datos."),
    ("Origen legacy", "Registro procedente del sistema historico; normalmente se conserva trazabilidad y centro."),
    ("Teixits-Mostres", "Circuito especifico de tejidos y muestras heredado del funcionamiento legacy."),
    ("Model", "Tipo de articulo/circuito historico del legacy, independiente de tejido o muestra."),
    ("Disposicion", "Documento operativo de preparacion o trabajo con cliente, color, piezas, kilos y acabador."),
    ("Parte de acabado", "Orden o parte productivo para planta/acabador."),
    ("Conteo ciego", "Inventario donde planta cuenta sin ver el stock esperado ni importes."),
    ("Intrastat", "Declaracion/agrupacion de operaciones intracomunitarias por pais y codigo arancelario."),
]


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, *, bold: bool = False, color: str | None = None) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_table(doc: Document, headers: list[str], rows: list[tuple[str, ...]], widths: list[Inches] | None = None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_text(cell, header, bold=True, color="FFFFFF")
        set_cell_shading(cell, "7D2D22")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        if widths:
            cell.width = widths[i]

    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], str(value))
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            if widths:
                cells[i].width = widths[i]

    doc.add_paragraph()
    return table


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_toc_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\\\o "1-3" \\\\h \\\\z \\\\u'
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(separate)
    run._r.append(end)


def add_footer(section) -> None:
    paragraph = section.footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("CoreFlow ERP SaaS - Manual de usuario - Documento interno")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(100, 100, 100)


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
    styles["Normal"].font.size = Pt(10.5)
    styles["Normal"].font.color.rgb = RGBColor(32, 43, 43)

    for name, size, color in [
        ("Title", 28, "7D2D22"),
        ("Heading 1", 18, "1F2A2A"),
        ("Heading 2", 14, "7D2D22"),
        ("Heading 3", 12, "1F2A2A"),
    ]:
        style = styles[name]
        style.font.name = "Aptos Display"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos Display")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)


def polish_paragraph_runs(paragraph) -> None:
    for run in paragraph.runs:
        if "/" in run.text:
            continue
        run.text = polish_text(run.text)


def polish_document(doc: Document) -> None:
    for paragraph in doc.paragraphs:
        polish_paragraph_runs(paragraph)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    polish_paragraph_runs(paragraph)


def add_cover(doc: Document) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("CoreFlow ERP SaaS")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(125, 45, 34)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("Manual de usuario")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Guia practica de navegacion, procesos y workflows operativos")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(86, 103, 103)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Version {VERSION} - {DOC_DATE}")

    notice = doc.add_paragraph()
    notice.alignment = WD_ALIGN_PARAGRAPH.CENTER
    notice.add_run("Documento interno para usuarios de la aplicacion. Incluye buenas practicas de uso diario y flujos principales.").italic = True
    doc.add_page_break()


def add_index(doc: Document) -> None:
    doc.add_heading("Indice", level=1)
    paragraph = doc.add_paragraph(
        "Este documento incluye una tabla de contenido automatica. En Microsoft Word, si los numeros de pagina no aparecen actualizados, usa Referencias > Actualizar tabla."
    )
    paragraph.runs[0].italic = True
    add_toc_field(doc.add_paragraph())

    doc.add_heading("Indice practico", level=2)
    for title, description in MANUAL_SECTIONS:
        paragraph = doc.add_paragraph(style="List Number")
        paragraph.add_run(title).bold = True
        paragraph.add_run(f": {description}")
    doc.add_page_break()


def add_area_intro(doc: Document, area_name: str) -> None:
    intros = {
        "CRM": "CRM concentra los datos de relacion comercial y operativa: quien compra, quien vende, quien transporta y que talleres intervienen. La calidad de estos maestros repercute directamente en documentos, estadisticas y trazabilidad.",
        "Articulos": "Articulos es el nucleo del catalogo textil y productivo. Conviene buscar siempre antes de crear, porque muchos registros proceden del legacy y pueden existir con codificacion historica.",
        "Base": "Base contiene parametros reutilizables en toda la aplicacion. Los cambios deben hacerse con cuidado porque pueden afectar a ventas, compras, produccion y reporting.",
        "Produccion": "Produccion permite controlar partes de acabado y ordenes vinculadas a muestras, models o disposiciones. Es el puente operativo entre ficha tecnica, trabajo de planta y stock.",
        "Ventas": "Ventas agrupa el circuito documental de cliente: pedido, albaran, pre-facturacion, factura, remesa e Intrastat. Las variantes Teixits-Mostres y Models mantienen la logica legacy donde corresponde.",
        "Compras": "Compras recoge el circuito de proveedor y los flujos especificos de hilos y fabricacion. Es importante revisar recepcion contra factura para detectar diferencias.",
        "Almacen": "Almacen muestra existencias, movimientos y operaciones fisicas. Es el area critica para validar stock real, traspasos, ajustes e inventarios de planta.",
    }
    doc.add_paragraph(intros[area_name])


def build_docx() -> None:
    DOCS.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    add_footer(section)
    configure_styles(doc)

    add_cover(doc)
    add_index(doc)

    doc.add_heading("1. Primeros pasos", level=1)
    doc.add_paragraph(
        "CoreFlow ERP SaaS centraliza el trabajo diario que antes estaba repartido en pantallas legacy: CRM, articulos, produccion, ventas, compras, almacen, listados, estadisticas y administracion. El objetivo del manual es que un usuario pueda encontrar la pantalla correcta, entender que se puede hacer y seguir los flujos habituales sin depender de conocimiento informal."
    )
    doc.add_heading("Antes de empezar", level=2)
    add_bullets(doc, [
        "Entra con tu usuario autorizado.",
        "Comprueba que la empresa activa es la correcta.",
        "Selecciona el idioma de pantalla si necesitas trabajar en castellano, catalan o ingles.",
        "Selecciona el idioma de impresiones/exportaciones si vas a generar PDFs, impresiones o CSV/XLSX.",
        "Elige tema visual: Actual, Azul moderno, Compacto o Ultracompacto.",
        "Despliega solo la categoria de menu que vas a usar para trabajar con mas foco.",
    ])

    doc.add_heading("2. Guia rapida para usuarios", level=1)
    doc.add_paragraph("Esta tabla responde a las dudas mas frecuentes de un usuario final: que pantalla abrir y que revisar antes de dar una tarea por buena.")
    add_table(doc, ["Necesidad", "Donde ir", "Que comprobar"], QUICK_RECIPES, [Inches(2.0), Inches(2.0), Inches(3.7)])

    doc.add_heading("3. Rutinas por rol", level=1)
    doc.add_paragraph("No todos los usuarios necesitan recorrer toda la aplicacion. Estas rutinas ayudan a cada perfil a empezar por lo que tiene mas impacto en su trabajo diario.")
    add_table(doc, ["Perfil", "Rutina recomendada"], ROLE_ROUTINES, [Inches(2.0), Inches(5.7)])

    doc.add_heading("4. Como leer una pantalla", level=1)
    doc.add_paragraph("Aunque cada modulo tenga campos propios, la mayoria de pantallas siguen la misma estructura. Entenderla reduce errores y acelera el aprendizaje.")
    add_table(doc, ["Zona", "Para que sirve"], SCREEN_ANATOMY, [Inches(2.0), Inches(5.7)])

    doc.add_heading("5. Capturas recomendadas", level=1)
    doc.add_paragraph("Para una version formativa completa, estas son las capturas que mas valor aportan. Conviene tomarlas con datos reales no sensibles, empresa correcta y tema visual acordado.")
    add_table(doc, ["Pantalla", "Que debe verse", "Para que sirve"], CAPTURE_PLAN, [Inches(2.1), Inches(3.1), Inches(2.5)])

    doc.add_heading("6. Estados, señales y datos criticos", level=1)
    doc.add_paragraph("Los estados y avisos son la forma mas rapida de saber si un registro requiere accion. Antes de cerrar un trabajo, revisar siempre estado, origen y datos criticos.")
    add_table(doc, ["Estado / señal", "Interpretacion practica"], STATUS_GUIDE, [Inches(2.1), Inches(5.6)])
    doc.add_heading("Datos que conviene validar siempre", level=2)
    add_table(doc, ["Dato", "Por que importa"], CRITICAL_DATA, [Inches(2.1), Inches(5.6)])

    doc.add_heading("7. Conceptos comunes", level=1)
    add_table(doc, ["Concepto", "Uso practico"], CONCEPTS, [Inches(1.8), Inches(5.9)])

    doc.add_heading("8. Mapa funcional", level=1)
    doc.add_paragraph("El mapa siguiente resume que se puede hacer en cada bloque principal y por donde se accede normalmente.")
    add_table(doc, ["Area", "Que permite", "Pantallas / rutas"], FUNCTIONAL_MAP, [Inches(1.4), Inches(3.9), Inches(2.5)])

    doc.add_heading("9. Workflows transversales", level=1)
    doc.add_paragraph("Estos procedimientos se repiten en muchas areas de la aplicacion. Si el usuario los domina, el resto de pantallas resultan mucho mas predecibles.")
    for title, steps in WORKFLOWS:
        doc.add_heading(title, level=2)
        add_numbered(doc, steps)

    chapter_no = 10
    for area_name, details in AREA_DETAILS.items():
        doc.add_heading(f"{chapter_no}. {area_name}", level=1)
        add_area_intro(doc, area_name)
        add_table(doc, ["Pantalla", "Uso"], details, [Inches(2.2), Inches(5.5)])
        doc.add_heading("Buenas practicas del area", level=2)
        add_bullets(doc, AREA_PRACTICES[area_name])
        chapter_no += 1

    doc.add_heading(f"{chapter_no}. Workflows operativos principales", level=1)
    doc.add_paragraph("Los siguientes flujos describen como resolver situaciones reales de trabajo. Estan redactados como procedimientos de usuario, no como documentacion tecnica.")
    for title, steps in SPECIFIC_WORKFLOWS:
        doc.add_heading(title, level=2)
        add_numbered(doc, steps)
    chapter_no += 1

    doc.add_heading(f"{chapter_no}. Procesos punta a punta", level=1)
    doc.add_paragraph("Estos recorridos muestran el flujo completo de trabajo, desde la entrada inicial hasta el resultado esperado. Son utiles para formar usuarios nuevos y para detectar donde se corta un circuito.")
    add_table(doc, ["Proceso", "Entrada", "Recorrido", "Resultado", "Validacion"], END_TO_END_FLOWS, [Inches(1.4), Inches(1.4), Inches(2.0), Inches(1.5), Inches(1.6)])
    chapter_no += 1

    doc.add_heading(f"{chapter_no}. Criterios de calidad", level=1)
    doc.add_paragraph("La regla practica es sencilla: antes de guardar, imprimir, exportar o cerrar, comprobar que el dato es completo, entendible y trazable. Esta tabla sirve como control minimo.")
    add_table(doc, ["Momento", "Que revisar", "Resultado esperado"], QUALITY_GATES, [Inches(1.7), Inches(3.1), Inches(2.9)])
    chapter_no += 1

    doc.add_heading(f"{chapter_no}. Resolucion de incidencias", level=1)
    doc.add_paragraph("Cuando algo no encaja, conviene seguir una respuesta ordenada: primero contexto, despues filtros, despues dato origen. Esta tabla evita perder tiempo mirando la pantalla equivocada.")
    add_table(doc, ["Sintoma", "Causa probable", "Primera accion"], INCIDENT_PLAYBOOK, [Inches(1.8), Inches(2.8), Inches(3.1)])
    chapter_no += 1

    doc.add_heading(f"{chapter_no}. Checklists de cierre", level=1)
    doc.add_paragraph("Los cierres no son solo contables: tambien sirven para dejar limpia la operativa de ventas, compras, produccion y almacen.")
    for title, items in CLOSING_CHECKLISTS.items():
        doc.add_heading(title, level=2)
        add_bullets(doc, items)
    chapter_no += 1

    doc.add_heading(f"{chapter_no}. Listados y estadisticas", level=1)
    doc.add_paragraph("Listados y Estadisticas permiten salir de la operacion pantalla a pantalla y leer el negocio por bloques. Son utiles para direccion, administracion y responsables de planta.")
    add_table(doc, ["Bloque", "Que revisar", "Accion recomendada"], [
        ("Ventas", "Mix cliente/articulo, concentracion, facturacion, pendiente de cobro y evolucion semanal.", "Revisar clientes principales, facturas pendientes y peso de pocos clientes sobre el total."),
        ("Compras", "Aging proveedor, recepcion vs factura, volumen por proveedor y evolucion.", "Cruzar recepciones con facturas antes de cerrar periodos."),
        ("Produccion", "Carga por acabador, estado y semana.", "Priorizar partes bloqueados o antiguos antes de lanzar trabajo nuevo."),
        ("Almacen", "Rotacion, antiguedad, cobertura de stock, movimientos y posiciones.", "Detectar stock parado, faltantes y diferencias recurrentes."),
    ], [Inches(1.5), Inches(3.3), Inches(3.0)])
    chapter_no += 1

    doc.add_heading(f"{chapter_no}. Alertas operativas", level=1)
    doc.add_paragraph("Alertas es una bandeja de señales. No sustituye al criterio del usuario, pero ayuda a no perder documentos o datos que requieren revision.")
    add_bullets(doc, [
        "Revisar alertas al iniciar la jornada.",
        "Resolver primero alertas que afecten a cliente, stock, facturacion o produccion viva.",
        "Abrir siempre el documento o maestro origen antes de corregir.",
        "Si una alerta no se puede resolver, documentar el motivo y revisarla mas adelante.",
    ])
    chapter_no += 1

    doc.add_heading(f"{chapter_no}. Administracion", level=1)
    doc.add_paragraph("Administracion solo aparece para usuarios con permisos. Permite gestionar tenant, plataforma, empresas, usuarios, acceso privado y configuracion general.")
    add_table(doc, ["Funcion", "Uso responsable"], [
        ("Empresas", "Mantener empresas y centros vinculados al tenant."),
        ("Usuarios", "Crear, invitar o revisar usuarios autorizados."),
        ("Cambio de contraseña", "Obligar o completar renovacion de credenciales cuando proceda."),
        ("Acceso privado", "Aprobar emails o claves de demo cuando la aplicacion esta en fase de pruebas."),
        ("Plataforma", "Uso reservado a administracion tecnica o plataforma."),
    ], [Inches(2.0), Inches(5.7)])
    chapter_no += 1

    doc.add_heading(f"{chapter_no}. Buenas practicas", level=1)
    add_bullets(doc, BEST_PRACTICES)
    chapter_no += 1

    doc.add_heading(f"{chapter_no}. Preguntas frecuentes", level=1)
    for question, answer in FAQ:
        doc.add_heading(question, level=2)
        doc.add_paragraph(answer)
    chapter_no += 1

    doc.add_heading(f"{chapter_no}. Glosario", level=1)
    add_table(doc, ["Termino", "Significado"], GLOSSARY, [Inches(2.0), Inches(5.7)])

    doc.add_heading("Checklist rapida de usuario", level=1)
    add_bullets(doc, [
        "Estoy en la empresa correcta.",
        "El idioma documental coincide con el destinatario del PDF o exportacion.",
        "He buscado antes de crear un registro nuevo.",
        "El documento tiene cliente/proveedor, fechas, importes y estado correctos.",
        "He revisado alertas o estadisticas si estoy cerrando periodo.",
        "He validado diferencias antes de cerrar inventario o regularizar stock.",
    ])

    doc.core_properties.title = "Manual de usuario CoreFlow ERP SaaS"
    doc.core_properties.subject = "Guia funcional de usuario"
    doc.core_properties.author = "CoreFlow / ERP SaaS"
    doc.core_properties.keywords = "ERP, manual, usuario, workflows, SaaS, legacy"
    polish_document(doc)
    doc.save(DOCX_PATH)


def build_markdown() -> None:
    lines: list[str] = [
        "# Manual de usuario CoreFlow ERP SaaS",
        "",
        f"Version {VERSION} - {DOC_DATE}",
        "",
        "## Indice",
    ]
    for idx, (title, description) in enumerate(MANUAL_SECTIONS, 1):
        lines.append(f"{idx}. **{title}**: {description}")

    lines.extend([
        "",
        "## 1. Primeros pasos",
        "CoreFlow ERP SaaS centraliza el trabajo diario que antes estaba repartido en pantallas legacy: CRM, articulos, produccion, ventas, compras, almacen, listados, estadisticas y administracion.",
        "",
        "### Antes de empezar",
    ])
    for item in [
        "Entra con tu usuario autorizado.",
        "Comprueba que la empresa activa es la correcta.",
        "Selecciona el idioma de pantalla.",
        "Selecciona el idioma de impresiones/exportaciones.",
        "Elige tema visual.",
        "Despliega solo la categoria de menu que vas a usar.",
    ]:
        lines.append(f"- {item}")

    lines.extend(["", "## 2. Guía rápida para usuarios"])
    for need, route, check in QUICK_RECIPES:
        lines.append(f"- **{need}**: ir a `{route}`. Comprobar: {check}")

    lines.extend(["", "## 3. Rutinas por rol"])
    for role, routine in ROLE_ROUTINES:
        lines.append(f"- **{role}**: {routine}")

    lines.extend(["", "## 4. Cómo leer una pantalla"])
    for zone, purpose in SCREEN_ANATOMY:
        lines.append(f"- **{zone}**: {purpose}")

    lines.extend(["", "## 5. Capturas recomendadas"])
    for screen, what, why in CAPTURE_PLAN:
        lines.append(f"- **{screen}**: debe verse {what} Sirve para: {why}")

    lines.extend(["", "## 6. Estados, señales y datos críticos"])
    for status, meaning in STATUS_GUIDE:
        lines.append(f"- **{status}**: {meaning}")
    lines.append("")
    lines.append("### Datos que conviene validar siempre")
    for data, reason in CRITICAL_DATA:
        lines.append(f"- **{data}**: {reason}")

    lines.extend(["", "## 7. Conceptos comunes"])
    for concept, usage in CONCEPTS:
        lines.append(f"- **{concept}**: {usage}")

    lines.extend(["", "## 8. Mapa funcional"])
    for area, what, routes in FUNCTIONAL_MAP:
        lines.append(f"- **{area}**: {what} Pantallas: `{routes}`")

    lines.extend(["", "## 9. Workflows transversales"])
    for title, steps in WORKFLOWS:
        lines.append(f"### {title}")
        for index, step in enumerate(steps, 1):
            lines.append(f"{index}. {step}")
        lines.append("")

    chapter_no = 10
    for area_name, details in AREA_DETAILS.items():
        lines.append(f"## {chapter_no}. {area_name}")
        for screen, usage in details:
            lines.append(f"- **{screen}**: {usage}")
        lines.append("")
        lines.append("### Buenas prácticas del área")
        for item in AREA_PRACTICES[area_name]:
            lines.append(f"- {item}")
        lines.append("")
        chapter_no += 1

    lines.append(f"## {chapter_no}. Workflows operativos principales")
    for title, steps in SPECIFIC_WORKFLOWS:
        lines.append(f"### {title}")
        for index, step in enumerate(steps, 1):
            lines.append(f"{index}. {step}")
        lines.append("")
    chapter_no += 1

    lines.append(f"## {chapter_no}. Procesos punta a punta")
    for process, source, route, result, validation in END_TO_END_FLOWS:
        lines.append(f"- **{process}**: entrada: {source}. Recorrido: {route}. Resultado: {result}. Validacion: {validation}")
    chapter_no += 1

    lines.append(f"## {chapter_no}. Criterios de calidad")
    for moment, review, expected in QUALITY_GATES:
        lines.append(f"- **{moment}**: revisar {review} Resultado esperado: {expected}")
    chapter_no += 1

    lines.append(f"## {chapter_no}. Resolucion de incidencias")
    for symptom, cause, action in INCIDENT_PLAYBOOK:
        lines.append(f"- **{symptom}**: causa probable: {cause} Primera accion: {action}")
    chapter_no += 1

    lines.append(f"## {chapter_no}. Checklists de cierre")
    for title, items in CLOSING_CHECKLISTS.items():
        lines.append(f"### {title}")
        for item in items:
            lines.append(f"- {item}")
    chapter_no += 1

    lines.append(f"## {chapter_no}. Listados y estadisticas")
    lines.append("Listados y Estadisticas permiten leer el negocio por bloques: ventas, compras, produccion y almacen.")
    chapter_no += 1

    lines.append(f"## {chapter_no}. Alertas operativas")
    for item in [
        "Revisar alertas al iniciar la jornada.",
        "Resolver primero alertas con impacto en cliente, stock, facturacion o produccion.",
        "Abrir el documento origen antes de corregir.",
    ]:
        lines.append(f"- {item}")
    chapter_no += 1

    lines.append(f"## {chapter_no}. Administracion")
    lines.append("Administracion aparece para usuarios con permisos y permite gestionar tenant, plataforma, empresas, usuarios y acceso privado.")
    chapter_no += 1

    lines.append(f"## {chapter_no}. Buenas practicas")
    for item in BEST_PRACTICES:
        lines.append(f"- {item}")
    chapter_no += 1

    lines.append(f"## {chapter_no}. Preguntas frecuentes")
    for question, answer in FAQ:
        lines.append(f"### {question}")
        lines.append(answer)
    chapter_no += 1

    lines.append(f"## {chapter_no}. Glosario")
    for term, meaning in GLOSSARY:
        lines.append(f"- **{term}**: {meaning}")

    polished_lines = [
        line if "Pantallas:" in line or "`/" in line else polish_text(line)
        for line in lines
    ]
    MD_PATH.write_text("\n".join(polished_lines) + "\n", encoding="utf-8")


if __name__ == "__main__":
    build_docx()
    build_markdown()
    print(DOCX_PATH)
    print(MD_PATH)
